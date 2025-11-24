from datetime import datetime
from docx import Document
from docx.shared import Inches
import dotenv
import matplotlib.pyplot as plt
import os
import pandas as pd
import requests
import textwrap

dotenv.load_dotenv()

# ---------- Settings ----------
CSV_PATH = "data/vendas_mercado.csv"
TEMPLATE_PATH = "input/template_relatorio.docx"
OUTPUT_DOCX = "output/relatorio_gerado.docx"
CHART_PNG = "output/grafico_vendas.png"
CHART_PNG_2 = "output/grafico_evolucao_vendas.png"
CHART_PNG_3 = "output/grafico_metodos_pagamento.png"
NAME = "Isabella"

# OpenRouter
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_MODEL = "meta-llama/llama-3.3-70b-instruct:free"


# ---------- Auxiliary functions ----------
def format_brl(value):
    """Format a number as Brazilian real currency. For example, 1234.56 -> R$ 1.234,56"""
    if pd.isna(value):
        return "R$ 0,00"
    return "R$ {:,.2f}".format(value).replace(",", "X").replace(".", ",").replace("X", ".")


def call_openrouter(prompt: str, max_tokens=600):
    """Call OpenRouter API to get a response for the given prompt."""
    if not OPENROUTER_API_KEY:
        raise RuntimeError("OPENROUTER_API_KEY não configurada.")

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": OPENROUTER_MODEL,
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "max_tokens": max_tokens,
        "temperature": 0.4
    }

    r = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=20)
    r.raise_for_status()

    return r.json()["choices"][0]["message"]["content"]


# ---------- Main script ----------
def main():
    df = pd.read_csv(CSV_PATH, parse_dates=["data"])

    if "receita_total" not in df.columns:
        df["receita_total"] = df["preco_unitario"] * df["quantidade"]

    # Calculate key metrics
    period_start = df["data"].min().date().strftime("%d/%m/%Y")
    period_end = df["data"].max().date().strftime("%d/%m/%Y")

    total_sales = float(df["receita_total"].sum())
    average_ticket = float(df["receita_total"].mean())

    category_revenue = (
        df.groupby("categoria", as_index=False)["receita_total"]
        .sum()
        .sort_values("receita_total", ascending=False)
    )

    top_category = category_revenue.iloc[0]["categoria"]
    top_category_revenue = category_revenue.iloc[0]["receita_total"]
    
    store_stats = df.groupby("loja").agg({
        "receita_total": "sum",
        "id_venda": "count",
        "margem": "sum"
    }).round(2)

    replacements = {
        "{{PERIODO_INICIO}}": period_start,
        "{{PERIODO_FIM}}": period_end,
        "{{TOTAL_VENDAS}}": format_brl(total_sales),
        "{{TOP_CATEGORIA}}": str(top_category),
        "{{FATURAMENTO_TOP_CATEGORIA}}": format_brl(top_category_revenue),
        "{{TICKET_MEDIO}}": format_brl(average_ticket),
        "{{NOME_DO_CANDIDATO}}": NAME,
        "{{DATA_GERACAO}}": datetime.now().strftime("%d/%m/%Y")
    }

    # Graph 1: Revenue by category
    plt.figure(figsize=(8, 5))
    category_revenue.plot(kind="bar", x="categoria", y="receita_total", legend=False)
    plt.title("Faturamento por categoria")
    plt.xlabel("Categoria")
    plt.ylabel("Receita (R$)")
    plt.tight_layout()
    plt.savefig(CHART_PNG)
    plt.close()
    
    # Graph 2: Sales evolution over time
    daily_sales = df.groupby(df["data"].dt.date)["receita_total"].sum().reset_index()
    daily_sales.columns = ["data", "receita_total"]
    
    plt.figure(figsize=(10, 5))
    plt.plot(daily_sales["data"], daily_sales["receita_total"], marker='o', linewidth=2, markersize=4)
    plt.title("Evolução de vendas ao longo do tempo")
    plt.xlabel("Data")
    plt.ylabel("Receita total (R$)")
    plt.xticks(rotation=45)
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    plt.savefig(CHART_PNG_2)
    plt.close()
    
    # Graph 3: Revenue by payment method
    payment_methods = df.groupby("metodo_pagamento")["receita_total"].sum().sort_values(ascending=False)
    
    plt.figure(figsize=(8, 8))
    colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#ff99cc']
    plt.pie(payment_methods.values, labels=payment_methods.index, autopct='%1.1f%%', 
            startangle=90, colors=colors[:len(payment_methods)])
    plt.title("Distribuição de receita por método de pagamento")
    plt.tight_layout()
    plt.savefig(CHART_PNG_3)
    plt.close()

    # Generate qualitative analysis via LLM
    prompt = textwrap.dedent(f"""
    Gere uma análise executiva de 5-10 linhas, em português, com tom claro.
    Dados:
    - Período: {period_start} a {period_end}
    - Total vendas: {format_brl(total_sales)}
    - Top categoria: {top_category} com {format_brl(top_category_revenue)}
    - Ticket médio: {format_brl(average_ticket)}

    Inclua insights práticos e um próximo passo de investigação.
    """)

    try:
        analysis = call_openrouter(prompt).strip()
    except Exception as e:
        analysis = f"(Erro ao gerar análise via OpenRouter: {e})"

    # Generate report
    doc = Document(TEMPLATE_PATH)

    for p in doc.paragraphs:
        for k, v in replacements.items():
            if k in p.text:
                p.text = p.text.replace(k, v)
        if "{{ANALISE_QUALITATIVA}}" in p.text:
            p.text = p.text.replace("{{ANALISE_QUALITATIVA}}", analysis)
        if "<<<INSERIR_GRAFICO_AQUI>>>" in p.text:
            p.text = ""
            p.add_run().add_picture(CHART_PNG, width=Inches(6))
        if "<<<INSERIR_GRAFICO_AQUI2>>>" in p.text:
            p.text = ""
            p.add_run().add_picture(CHART_PNG_2, width=Inches(6))
        if "<<<INSERIR_GRAFICO_AQUI3>>>" in p.text:
            p.text = ""
            p.add_run().add_picture(CHART_PNG_3, width=Inches(6))
        if "<<<INSERIR_TABELA_AQUI>>>" in p.text:
            p.text = ""
            table = doc.add_table(rows=1 + len(store_stats), cols=4)
            table.style = 'Light Grid Accent 1'
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "Loja"
            header_cells[1].text = "Faturamento Total"
            header_cells[2].text = "Nº Transações"
            header_cells[3].text = "Margem Total"
            
            for idx, store in enumerate(store_stats.index, start=1):
                row_cells = table.rows[idx].cells
                row_cells[0].text = store
                row_cells[1].text = format_brl(store_stats.loc[store, "receita_total"])
                row_cells[2].text = str(int(store_stats.loc[store, "id_venda"]))
                row_cells[3].text = format_brl(store_stats.loc[store, "margem"])
            
            p._element.addnext(table._element)

    doc.save(OUTPUT_DOCX)
    print(f"Relatório gerado com sucesso: {OUTPUT_DOCX}")



if __name__ == "__main__":
    main()
    